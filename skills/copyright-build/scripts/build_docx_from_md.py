#!/usr/bin/env python3
"""Build final DOCX/TXT files from confirmed Markdown drafts."""

from __future__ import annotations

import argparse
import html
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from common import ensure_dir, read_json, safe_filename

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


BLACK_RGB = "000000"

# 复用竞赛 Node 公式链，把 LaTeX 转为 Word 可编辑的 OMML 矢量公式；不可用时降级为原文文本。
_omml = None
try:
    import sys as _sys
    import os as _os
    # 优先用平台注入的引擎绝对路径（工作区跑脚本时向上找不到 tools/docx-cn-engine，
    # 后端把主仓库引擎位置通过 DOCX_CN_ENGINE_DIR 传进来）；否则再向上逐级搜索兜底。
    _cand_dirs = []
    _env_engine = _os.environ.get("DOCX_CN_ENGINE_DIR", "").strip()
    if _env_engine:
        _cand_dirs.append(Path(_env_engine))
    _cur = Path(__file__).resolve().parent
    for _ in range(8):
        _cand_dirs.append(_cur / "tools" / "docx-cn-engine")
        if _cur.parent == _cur:
            break
        _cur = _cur.parent
    for _cand in _cand_dirs:
        if (_cand / "omml_helper.py").is_file():
            if str(_cand) not in _sys.path:
                _sys.path.insert(0, str(_cand))
            import omml_helper as _omml  # type: ignore
            break
except Exception:
    _omml = None

_OMML_CACHE: dict[tuple[str, bool], "str | None"] = {}
_INLINE_MATH_BARE_RE = re.compile(r"(?<!\$)\$(?!\$)((?:\\.|[^$\n])+?)\$(?!\$)")
_INLINE_MATH_PAREN_BARE_RE = re.compile(r"\\\(((?:\\.|[^\\\n])+?)\\\)")


def _omml_available() -> bool:
    return _omml is not None and _omml.omml_available()


def _omml_lookup(latex: str, display: bool):
    key = (latex.strip(), bool(display))
    if key in _OMML_CACHE:
        return _OMML_CACHE[key]
    if not _omml_available():
        _OMML_CACHE[key] = None
        return None
    res = _omml.render_omml_batch([(key[0], key[1])])
    val = res[0] if res else None
    _OMML_CACHE[key] = val
    return val


def _prewarm_omml(md_text: str) -> None:
    """扫描全文收集 LaTeX 公式，一次性批量转 OMML 缓存（跳过围栏代码块）。"""
    _OMML_CACHE.clear()
    if not _omml_available():
        return
    items: list[tuple[str, bool]] = []
    seen: set[tuple[str, bool]] = set()

    def _add(latex: str, display: bool):
        k = (latex.strip(), bool(display))
        if k[0] and k not in seen:
            seen.add(k)
            items.append(k)

    lines = md_text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if s.startswith("```"):
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue
        if s == "$$":
            i += 1
            body: list[str] = []
            while i < n and lines[i].strip() != "$$":
                body.append(lines[i])
                i += 1
            i += 1
            _add("\n".join(body), True)
            continue
        if s == "\\[":
            i += 1
            body = []
            while i < n and lines[i].strip() != "\\]":
                body.append(lines[i])
                i += 1
            i += 1
            _add("\n".join(body), True)
            continue
        if s.startswith("$$") and s.endswith("$$") and len(s) > 4:
            _add(s[2:-2], True)
            i += 1
            continue
        # 单行 \[ ... \]（块级公式写在一行）
        if s.startswith("\\[") and s.endswith("\\]") and len(s) > 4:
            _add(s[2:-2], True)
            i += 1
            continue
        for m in _INLINE_MATH_BARE_RE.finditer(lines[i]):
            _add(m.group(1), False)
        for m in _INLINE_MATH_PAREN_BARE_RE.finditer(lines[i]):
            _add(m.group(1), False)
        i += 1

    if not items:
        return
    results = _omml.render_omml_batch(items)
    for (latex, disp), omml in zip(items, results):
        _OMML_CACHE[(latex, disp)] = omml


def strip_markdown_links(text: str) -> str:
    text = re.sub(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"<(https?://[^>]+)>", r"\1", text)
    return text


def parse_application_lines(md_path: Path) -> tuple[list[str], list[str]]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    fields = [line.strip() for line in lines if line.strip().startswith("➤")]
    warnings = [line for line in fields if "待用户确认" in line]
    return fields, warnings


def parse_application_field(md_path: Path, field_name: str) -> str:
    if not md_path.exists():
        return ""
    prefix = f"➤{field_name}："
    for line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith(prefix):
            return stripped[len(prefix) :].strip()
    return ""


def application_version(draft_dir: Path) -> str:
    version = parse_application_field(draft_dir / "申请表信息.md", "版本号")
    if "待用户确认" in version:
        return ""
    return version


def application_software_name(draft_dir: Path) -> str:
    name = parse_application_field(draft_dir / "申请表信息.md", "软件全称")
    if "待用户确认" in name:
        return ""
    return name


def write_application_txt(
    draft_dir: Path, out_dir: Path, total_pages: int | None = None
) -> tuple[Path | None, list[str]]:
    md_path = draft_dir / "申请表信息.md"
    if not md_path.exists():
        return None, ["缺少草稿/申请表信息.md"]
    fields, warnings = parse_application_lines(md_path)
    # ⛔ 回填页数：代码 docx 折叠空行+重排后实际页数可能与 draft 声称的不同，
    #    用重排后的真实总页数覆盖 ➤页数，保证申请表与实际 PDF 一致（total_pages>0 才回填）。
    if total_pages and total_pages > 0:
        backfilled = False
        for i, ln in enumerate(fields):
            # 匹配"页数"字段行（全角/半角冒号，行首可有 ➤ 或空白）
            if re.match(r"^\s*[➤\-\*]?\s*页\s*数\s*[:：]", ln):
                prefix = ln.split("：", 1)[0] if "：" in ln else ln.split(":", 1)[0]
                sep = "：" if "：" in ln else ":"
                fields[i] = f"{prefix}{sep}{total_pages}"
                backfilled = True
                break
        if not backfilled:
            warnings.append(f"申请表未找到“页数”字段，未能回填真实页数（实际 {total_pages} 页）")
    out_path = out_dir / "申请表信息.txt"
    out_path.write_text("\n".join(fields) + "\n", encoding="utf-8")
    return out_path, warnings


def read_json_if_exists(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return read_json(path)


def confirmation_issues(workdir: Path) -> list[str]:
    draft_dir = workdir / "草稿"
    issues: list[str] = []

    business = read_json_if_exists(draft_dir / "业务理解.json")
    if not business or not business.get("user_confirmed"):
        issues.append("业务理解尚未确认：请确认 草稿/业务理解.md 后记录 `business` 门禁")

    selection = read_json_if_exists(draft_dir / "代码文件选择.json")
    if not selection or not selection.get("user_confirmed"):
        issues.append("代码文件选择尚未确认：请确认 草稿/代码文件选择.json 后记录 `code-selection` 门禁")

    screenshot = read_json_if_exists(workdir / "截图方式确认.json")
    if not screenshot.get("screenshot_method_confirmed"):
        issues.append("截图方式尚未确认：请选择截图方式后记录 `screenshot-method` 门禁")

    app_md = draft_dir / "申请表信息.md"
    if app_md.exists():
        _, warnings = parse_application_lines(app_md)
        if warnings:
            issues.append("申请表信息仍包含“待用户确认”字段")
    else:
        issues.append("缺少 草稿/申请表信息.md")

    app_confirmation = read_json_if_exists(draft_dir / "申请表字段确认.json")
    if not app_confirmation.get("application_fields_confirmed"):
        issues.append("申请表字段尚未确认：请补全字段后记录 `application-fields` 门禁")

    markdown_confirmation = read_json_if_exists(draft_dir / "最终生成确认.json")
    if not markdown_confirmation.get("markdown_confirmed"):
        issues.append("Markdown 草稿尚未最终确认：请确认全部草稿后记录 `markdown` 门禁")

    return issues


def parse_code_pages(md_path: Path) -> list[tuple[int, list[str]]]:
    pages: list[tuple[int, list[str]]] = []
    current_no: int | None = None
    current_lines: list[str] = []
    in_fence = False

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        page_match = re.match(r"^##\s+第\s*(\d+)\s*页", raw.strip())
        if page_match:
            if current_no is not None:
                pages.append((current_no, current_lines))
            current_no = int(page_match.group(1))
            current_lines = []
            in_fence = False
            continue
        if raw.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if current_no is not None and in_fence:
            current_lines.append(raw)

    if current_no is not None:
        pages.append((current_no, current_lines))
    return pages


# 每页代码行数：控在这个范围内重排，消灭"随机 1-2 行短页"，并让每页尽量填满。
CODE_LINES_PER_PAGE = 48


def reflow_code_pages(
    pages: list[tuple[int, list[str]]], start_page: int = 1
) -> list[tuple[int, list[str]]]:
    """把解析出的页重排，解决两类真实翻车：
      ① 随机短页（某页只有 1-2 行代码，下面大片空白，页数随机）——根因是 draft
         切页不均、build 无条件信任页边界。这里把所有行汇成一条流、按定长重切，
         每页都填满（末页除外），短页消失。
      ② 空行过多（连续空两行/多行原样占版面，有效代码少）——把 ≥2 个连续空行
         折叠成 1 个，并去掉每页首尾空行，版面利用率提升。
    页码从 start_page 连续重排，返回 (页码, 行列表) 列表。跨文件时由调用方把上一
    文件的末页+1 传进来，保证前30页/后30页连续编号。
    """
    # 1. 汇成一条行流（忽略原页边界）
    flat: list[str] = []
    for _no, lines in pages:
        flat.extend(lines)
    # 2. 折叠连续空行：≥2 个连续空行 → 1 个；同时丢弃流首尾的空行
    collapsed: list[str] = []
    blank_run = 0
    for ln in flat:
        if ln.strip() == "":
            blank_run += 1
            if blank_run >= 2:
                continue  # 第 2 个及以后的连续空行丢弃
            collapsed.append("")
        else:
            blank_run = 0
            collapsed.append(ln.rstrip())  # 去行尾空白，避免尾部空格占位
    # 去掉整体首尾空行
    while collapsed and collapsed[0] == "":
        collapsed.pop(0)
    while collapsed and collapsed[-1] == "":
        collapsed.pop()
    if not collapsed:
        return [(start_page, [])]
    # 3. 按定长重切；切页时避免把页首留成空行
    result: list[tuple[int, list[str]]] = []
    page_no = start_page
    i = 0
    n = len(collapsed)
    while i < n:
        chunk = collapsed[i:i + CODE_LINES_PER_PAGE]
        i += CODE_LINES_PER_PAGE
        # 下一页页首若是空行，顺手并入本页尾（不新增行数上限太多，最多 +1）
        while i < n and collapsed[i] == "" and len(chunk) < CODE_LINES_PER_PAGE + 1:
            chunk.append("")
            i += 1
        # 去掉本页首尾空行（重切后页边界可能落在空行上）
        while chunk and chunk[0] == "":
            chunk.pop(0)
        while chunk and chunk[-1] == "":
            chunk.pop()
        if chunk:
            result.append((page_no, chunk))
            page_no += 1
    return result if result else [(start_page, [])]


def set_run_font(run: Any, name: str, size_pt: float) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


_HEADING_SIZE_PT = {1: 15, 2: 14, 3: 12, 4: 11}


def add_manual_heading(document: Any, level: int, text: str) -> None:
    """对标竞赛模板：手工构建标题段落，显式黑色加粗 + 明确字号，
    避免 Word 内置"标题 N"样式的主题蓝色与大纲方块符号。"""
    lvl = min(max(level, 1), 4)
    size_pt = _HEADING_SIZE_PT[lvl]
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if lvl <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "SimHei" if lvl <= 2 else "SimSun", size_pt)
    run.font.bold = True


_INLINE_MD_RE = re.compile(r"(\*\*[^*\n]+?\*\*|`[^`\n]+?`)")


def strip_inline_markers(text: str) -> str:
    """去掉行内 **粗体** / `代码` 标记，只保留可读文本（用于标题等不分 run 的场景）。"""
    text = re.sub(r"\*\*([^*\n]+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`\n]+?)`", r"\1", text)
    return text


def add_inline_runs(paragraph: Any, text: str, size_pt: float = 10.5) -> None:
    """把一段文字按 **粗体** / `代码` / 行内公式 拆成多个 run 写入段落，其余按正文宋体。"""
    # 收集 token：粗体/代码（原有）+ 行内公式（OMML 可用时）
    spans: list[tuple[int, int, str, str]] = []
    for m in _INLINE_MD_RE.finditer(text):
        spans.append((m.start(), m.end(), "md", m.group(1)))
    if _omml_available():
        taken = [(s, e) for s, e, _, _ in spans]

        def _overlap(a: int, b: int) -> bool:
            return any(not (b <= s or a >= e) for s, e in taken)

        for m in _INLINE_MATH_BARE_RE.finditer(text):
            if not _overlap(m.start(), m.end()):
                spans.append((m.start(), m.end(), "math", m.group(1)))
                taken.append((m.start(), m.end()))
        for m in _INLINE_MATH_PAREN_BARE_RE.finditer(text):
            if not _overlap(m.start(), m.end()):
                spans.append((m.start(), m.end(), "math", m.group(1)))
                taken.append((m.start(), m.end()))

    spans.sort(key=lambda t: t[0])
    pos = 0
    for start, end, kind, token in spans:
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            set_run_font(run, "SimSun", size_pt)
        if kind == "md":
            if token.startswith("**"):
                run = paragraph.add_run(token[2:-2])
                set_run_font(run, "SimSun", size_pt)
                run.font.bold = True
            else:
                run = paragraph.add_run(token[1:-1])
                set_run_font(run, "Consolas", max(size_pt - 1.5, 8))
        else:  # math：转 OMML，失败则原文
            omml_xml = _omml_lookup(token, False)
            if not (omml_xml and _omml is not None and _omml.append_inline_omml(paragraph, omml_xml)):
                run = paragraph.add_run(f"${token}$")
                set_run_font(run, "SimSun", size_pt)
        pos = end
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, "SimSun", size_pt)


def set_normal_font(document: Any, name: str = "SimSun", size_pt: float = 10.5) -> None:
    style = document.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    try:
        style.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def set_style_black(document: Any) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        try:
            document.styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass


def force_black_document(document: Any) -> None:
    set_style_black(document)
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            except Exception:
                                pass


def configure_a4(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)


def configure_code_a4(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def add_page_field(paragraph: Any) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate, result, end):
        run = paragraph.add_run()
        run._r.append(element)
        set_run_font(run, "SimSun", 8)


def set_code_header(document: Any, software_name: str, version: str) -> None:
    section = document.sections[0]
    section.header.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].text = "" if header.paragraphs else None

    # Build a two-column header: software name on the left, page number on the right.
    table = header.add_table(rows=1, cols=2, width=Cm(17.5))
    table.autofit = True
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    left_para.paragraph_format.line_spacing = Pt(12)
    left_run = left_para.add_run(f"{software_name} {version}")
    set_run_font(left_run, "SimSun", 8)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    right_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    right_para.paragraph_format.line_spacing = Pt(12)
    prefix = right_para.add_run("第 ")
    set_run_font(prefix, "SimSun", 8)
    add_page_field(right_para)
    suffix = right_para.add_run(" 页")
    set_run_font(suffix, "SimSun", 8)

    # Remove borders from the header table
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_borders.append(border)
        tc_pr.append(tc_borders)


def build_code_docx_python(
    md_path: Path, out_path: Path, software_name: str, version: str, start_page: int = 1
) -> int:
    """生成代码 docx，返回本文件占用的页数（供跨文件连续编号）。"""
    raw_pages = parse_code_pages(md_path)
    if not raw_pages:
        raise RuntimeError(f"No code pages parsed from {md_path}")
    # 折叠空行 + 定长重排：消灭随机短页/多余空行，页码从 start_page 连续
    pages = reflow_code_pages(raw_pages, start_page=start_page)

    document = Document()
    configure_code_a4(document)
    set_normal_font(document, "Consolas", 7.2)
    set_style_black(document)
    set_code_header(document, software_name, version)

    # 页码从 start_page 起（前30页=1，后30页接前一文件末页+1，实现连续编号）
    start_page_no = pages[0][0] if pages else 1
    if start_page_no != 1:
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), str(start_page_no))
        document.sections[0]._sectPr.append(pg_num_type)

    for index, (page_no, lines) in enumerate(pages):
        for line in lines:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(14)
            run = p.add_run(line if line else " ")
            set_run_font(run, "Consolas", 7.2)
        if index != len(pages) - 1:
            # 嵌入式分页符：避免 add_page_break() 产生多余空段落导致空白页
            run = p.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)

    force_black_document(document)
    document.save(out_path)
    return len(pages)


def paragraph_xml(text: str, font: str = "SimSun", size_half_points: int = 21, align: str | None = None, line_twips: int = 240) -> str:
    align_xml = f'<w:jc w:val="{align}"/>' if align else ""
    escaped = html.escape(text)
    return (
        "<w:p>"
        f"<w:pPr>{align_xml}<w:spacing w:after=\"0\" w:line=\"{line_twips}\" w:lineRule=\"exact\"/></w:pPr>"
        "<w:r>"
        f"<w:rPr><w:rFonts w:ascii=\"{font}\" w:hAnsi=\"{font}\" w:eastAsia=\"{font}\"/>"
        f"<w:color w:val=\"{BLACK_RGB}\"/>"
        f"<w:sz w:val=\"{size_half_points}\"/><w:szCs w:val=\"{size_half_points}\"/></w:rPr>"
        f"<w:t xml:space=\"preserve\">{escaped}</w:t>"
        "</w:r>"
        "</w:p>"
    )


def page_break_xml() -> str:
    return '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'


def page_field_runs_xml() -> str:
    return (
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:t>1</w:t></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="end"/></w:r>'
    )


def header_xml(header_text: str) -> str:
    """Build a two-column header: software name left, page number right."""
    escaped = html.escape(header_text)
    # Use a borderless table for left/right alignment in header
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:tbl>
    <w:tblPr>
      <w:tblW w:w="5000" w:type="pct"/>
      <w:tblBorders>
        <w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>
      </w:tblBorders>
    </w:tblPr>
    <w:tr>
      <w:tc>
        <w:p>
          <w:pPr><w:jc w:val="left"/><w:spacing w:after="0" w:line="240" w:lineRule="exact"/></w:pPr>
          <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve">{escaped}</w:t></w:r>
        </w:p>
      </w:tc>
      <w:tc>
        <w:p>
          <w:pPr><w:jc w:val="right"/><w:spacing w:after="0" w:line="240" w:lineRule="exact"/></w:pPr>
          <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve">第 </w:t></w:r>
          {page_field_runs_xml()}
          <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve"> 页</w:t></w:r>
        </w:p>
      </w:tc>
    </w:tr>
  </w:tbl>
</w:hdr>"""


def minimal_docx(out_path: Path, body_xml: str, header_text: str | None = None, start_page: int = 1) -> None:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    header_rel = (
        '<Relationship Id="rIdHeader1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="header1.xml"/>'
        if header_text
        else ""
    )
    doc_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{header_rel}</Relationships>"""
    styles = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="21"/></w:rPr>
  </w:style>
</w:styles>"""
    document = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    {body_xml}
      <w:sectPr>
        {'<w:headerReference w:type="default" r:id="rIdHeader1"/>' if header_text else ''}
        {'<w:pgNumType w:start="' + str(start_page) + '"/>' if start_page != 1 else ''}
        <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1134" w:right="1134" w:bottom="1134" w:left="1418" w:header="283" w:footer="283" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>"""
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/styles.xml", styles)
        zf.writestr("word/document.xml", document)
        if header_text:
            zf.writestr("word/header1.xml", header_xml(header_text))


def force_black_xml(xml: str) -> str:
    xml = re.sub(r"<w:hyperlink\b[^>]*>", "", xml)
    xml = xml.replace("</w:hyperlink>", "")
    xml = re.sub(r"<w:color\b[^>]*/>", f'<w:color w:val="{BLACK_RGB}"/>', xml)

    def ensure_rpr_color(match: re.Match[str]) -> str:
        value = match.group(0)
        if "<w:color" in value:
            return value
        return value.replace("</w:rPr>", f'<w:color w:val="{BLACK_RGB}"/></w:rPr>')

    xml = re.sub(r"<w:rPr\b[^>]*>.*?</w:rPr>", ensure_rpr_color, xml, flags=re.S)
    xml = re.sub(r"<w:r>(?!<w:rPr>)", f'<w:r><w:rPr><w:color w:val="{BLACK_RGB}"/></w:rPr>', xml)
    return xml


def normalize_docx_text_color(docx_path: Path) -> None:
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    color_xml_parts = (
        "word/document.xml",
        "word/styles.xml",
        "word/numbering.xml",
        "word/header",
        "word/footer",
    )
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.endswith(".xml") and item.filename.startswith(color_xml_parts):
                text = data.decode("utf-8")
                data = force_black_xml(text).encode("utf-8")
            elif item.filename.endswith(".rels"):
                text = data.decode("utf-8", errors="ignore")
                if "hyperlink" in text:
                    text = re.sub(r'\s*<Relationship\b[^>]*Type="[^"]*/hyperlink"[^>]*/>', "", text)
                    data = text.encode("utf-8")
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def next_header_part(names: set[str]) -> tuple[str, str]:
    index = 1
    while f"word/header{index}.xml" in names:
        index += 1
    return f"word/header{index}.xml", f"header{index}.xml"


def unique_relationship_id(rels_xml: str, base: str = "rIdManualHeader") -> str:
    if f'Id="{base}"' not in rels_xml:
        return base
    index = 2
    while f'Id="{base}{index}"' in rels_xml:
        index += 1
    return f"{base}{index}"


def add_header_to_existing_docx(docx_path: Path, header_text: str) -> None:
    """Add the same two-column header used by code materials to an existing DOCX."""
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as src:
        names = set(src.namelist())
        header_part, header_target = next_header_part(names)
        rels_xml = src.read("word/_rels/document.xml.rels").decode("utf-8")
        rel_id = unique_relationship_id(rels_xml)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    override = (
                        f'<Override PartName="/{header_part}" '
                        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>'
                    )
                    if f'PartName="/{header_part}"' not in text:
                        text = text.replace("</Types>", f"{override}</Types>")
                    data = text.encode("utf-8")
                elif item.filename == "word/_rels/document.xml.rels":
                    text = data.decode("utf-8")
                    relationship = (
                        f'<Relationship Id="{rel_id}" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" '
                        f'Target="{header_target}"/>'
                    )
                    text = text.replace("</Relationships>", f"{relationship}</Relationships>")
                    data = text.encode("utf-8")
                elif item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    if "xmlns:r=" not in text:
                        text = text.replace(
                            "<w:document ",
                            '<w:document xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" ',
                            1,
                        )
                    header_ref = f'<w:headerReference w:type="default" r:id="{rel_id}"/>'
                    if "<w:headerReference" in text:
                        text = re.sub(r"<w:headerReference\b[^>]*/>", header_ref, text, count=1)
                    else:
                        text = re.sub(r"(<w:sectPr\b[^>]*>)", rf"\1{header_ref}", text, count=1)
                    data = text.encode("utf-8")
                dst.writestr(item, data)
            dst.writestr(header_part, header_xml(header_text))
    tmp_path.replace(docx_path)


def build_code_docx_ooxml(
    md_path: Path, out_path: Path, software_name: str, version: str, start_page: int = 1
) -> int:
    """OOXML 兜底版：同样折叠空行+定长重排，返回本文件页数。"""
    raw_pages = parse_code_pages(md_path)
    if not raw_pages:
        raise RuntimeError(f"No code pages parsed from {md_path}")
    pages = reflow_code_pages(raw_pages, start_page=start_page)
    start_page_no = pages[0][0] if pages else 1
    body: list[str] = []
    for index, (page_no, lines) in enumerate(pages):
        for line in lines:
            body.append(paragraph_xml(line if line else " ", font="Consolas", size_half_points=14, line_twips=280))
        if index != len(pages) - 1:
            # 嵌入式分页符：嵌入最后一段的 run 避免多余空段落
            last = body.pop()
            last = last.replace('</w:r></w:p>', '<w:br w:type="page"/></w:r></w:p>')
            body.append(last)
    minimal_docx(out_path, "\n".join(body), header_text=f"{software_name} {version}", start_page=start_page_no)
    return len(pages)


def add_markdown_table(document: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    def fill_cell(cell: Any, text: str) -> None:
        cell.text = ""
        add_inline_runs(cell.paragraphs[0], strip_markdown_links(text), 10.5)

    for idx, text in enumerate(rows[0]):
        fill_cell(table.rows[0].cells[idx], text)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(cells)]):
            fill_cell(cells[idx], text)


def parse_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def image_pixel_size(path: Path) -> tuple[int, int] | None:
    """读取常见位图宽高（像素），失败返回 None。不依赖 Pillow。"""
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if len(raw) >= 24 and raw.startswith(b"\x89PNG\r\n\x1a\n") and raw[12:16] == b"IHDR":
        w = int.from_bytes(raw[16:20], "big")
        h = int.from_bytes(raw[20:24], "big")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 10 and raw[:3] == b"GIF" and raw[3:6] in (b"87a", b"89a"):
        w = int.from_bytes(raw[6:8], "little")
        h = int.from_bytes(raw[8:10], "little")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 4 and raw.startswith(b"\xff\xd8"):
        i = 2
        n = len(raw)
        while i < n:
            if raw[i] != 0xFF:
                i += 1
                continue
            i += 1
            while i < n and raw[i] == 0xFF:
                i += 1
            if i >= n:
                break
            marker = raw[i]
            i += 1
            if marker in (0xD8, 0xD9):
                continue
            if marker == 0xDA:
                break
            if 0xD0 <= marker <= 0xD7:
                continue
            if i + 2 > n:
                break
            seg_len = int.from_bytes(raw[i : i + 2], "big")
            if seg_len < 2:
                break
            i += 2
            if marker in (0xC0, 0xC1, 0xC2) and i + 5 <= n:
                h = int.from_bytes(raw[i + 1 : i + 3], "big")
                w = int.from_bytes(raw[i + 3 : i + 5], "big")
                if w > 0 and h > 0:
                    return w, h
            i += seg_len - 2
    return None


def add_image(
    document: Any,
    image_path: Path,
    caption_desc: str = "",
    fig_counter: list[int] | None = None,
) -> None:
    if not image_path.exists():
        p = document.add_paragraph()
        run = p.add_run(f"[截图缺失：{image_path}]")
        set_run_font(run, "SimSun", 10.5)
        return
    # 尺寸规范化：正常图统一目标宽 12cm（14cm 在手册页里偏大、挤走正文导致图文失衡，
    # 用户反馈问题②；12cm 半栏偏大、图文更平衡）；低像素图防糊；瘦长图防溢出
    target_width_cm = 12.0
    max_height_cm = 18.0
    actual_width_cm = target_width_cm
    dims = image_pixel_size(image_path)
    if dims:
        img_w_px, img_h_px = dims
        eff_dpi = img_w_px / (target_width_cm / 2.54)
        if eff_dpi < 150:
            native_cm = img_w_px / 96 * 2.54
            actual_width_cm = min(native_cm, target_width_cm)
        if img_w_px > 0:
            aspect = img_h_px / img_w_px
            height_cm = actual_width_cm * aspect
            if height_cm > max_height_cm:
                actual_width_cm = max_height_cm / aspect
    try:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(actual_width_cm))
    except Exception:
        p = document.add_paragraph()
        run = p.add_run(f"[截图无法插入：{image_path}]")
        set_run_font(run, "SimSun", 10.5)
        return
    # 图下方居中题注："图 N 描述"
    if fig_counter is not None:
        fig_counter[0] += 1
        desc = (caption_desc or "").strip() or "界面截图"
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Pt(0)
        num_run = cap.add_run(f"图 {fig_counter[0]}")
        set_run_font(num_run, "SimSun", 9)
        num_run.font.bold = True
        desc_run = cap.add_run(f" {desc}")
        set_run_font(desc_run, "SimSun", 9)


def build_manual_docx_python(md_path: Path, out_path: Path, base_dir: Path, software_name: str, version: str) -> None:
    document = Document()
    configure_a4(document)
    set_normal_font(document, "SimSun", 10.5)
    set_style_black(document)
    set_code_header(document, software_name, version)
    md_text = md_path.read_text(encoding="utf-8")
    _prewarm_omml(md_text)  # 预热公式 OMML，一次性批量转换
    lines = md_text.splitlines()
    table_buf: list[list[str]] = []
    in_fence = False
    fig_no = [0]
    list_counter = 0  # 连续列表项计数（遇非列表内容归零，用于中文编号「（N）」）
    block_math_buf: list[str] | None = None  # 收集块级公式 $$ / \[ 内多行

    def flush_block_math(delim_close: str) -> None:
        nonlocal block_math_buf, list_counter
        if block_math_buf is None:
            return
        latex = "\n".join(block_math_buf).strip()
        block_math_buf = None
        list_counter = 0
        omml_xml = _omml_lookup(latex, True) if latex else None
        if omml_xml and _omml is not None and _omml.add_display_omml(document, omml_xml):
            return
        # 降级：公式原文居中一行
        if latex:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(latex)
            set_run_font(run, "SimSun", 12)

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            data = [row for row in table_buf if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)]
            add_markdown_table(document, data)
            table_buf = []

    for line in lines:
        stripped = line.strip()
        stripped = strip_markdown_links(stripped)
        if stripped.startswith("```"):
            flush_table()
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        # 块级公式：$$…$$ 或 \[…\]（可跨多行）
        if block_math_buf is not None:
            if stripped == "$$" or stripped == "\\]":
                flush_block_math(stripped)
            else:
                block_math_buf.append(line)
            continue
        if stripped == "$$" or stripped == "\\[":
            flush_table()
            block_math_buf = []
            continue
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            flush_table()
            block_math_buf = [stripped[2:-2]]
            flush_block_math("$$")
            continue
        # 单行块级公式：\[ ... \] 写在同一行（否则会当普通文字漏转）
        if stripped.startswith("\\[") and stripped.endswith("\\]") and len(stripped) > 4:
            flush_table()
            block_math_buf = [stripped[2:-2].strip()]
            flush_block_math("\\]")
            continue
        if stripped.startswith("<!--") and "截图" in stripped:
            stripped = "【截图预留：请在此处插入当前功能页面或操作结果截图。】"
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buf.append(parse_table_line(stripped))
            list_counter = 0
            continue
        flush_table()
        if not stripped:
            continue
        image_match = re.search(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            list_counter = 0
            add_image(
                document,
                (base_dir / image_match.group(2)).resolve(),
                caption_desc=image_match.group(1),
                fig_counter=fig_no,
            )
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            list_counter = 0
            level = min(len(heading.group(1)), 4)
            add_manual_heading(document, level, strip_inline_markers(heading.group(2)))
            continue
        # 无序 / 有序列表：统一转中文编号「（N）」，连续项递增，不用内置 • 圆点样式
        list_match = re.match(r"^(?:[-*+]|\d+\.)\s+(.+)$", stripped)
        if list_match:
            list_counter += 1
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.first_line_indent = Pt(2 * 12)
            prefix_run = p.add_run(f"（{list_counter}）")
            set_run_font(prefix_run, "SimSun", 12)
            add_inline_runs(p, list_match.group(1), 12)
            continue
        list_counter = 0
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.first_line_indent = Pt(2 * 12)
        add_inline_runs(p, stripped, 12)
    flush_table()
    force_black_document(document)
    document.save(out_path)


def pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def build_with_pandoc(md_path: Path, out_path: Path, code_mode: bool = False) -> None:
    if not pandoc_available():
        raise RuntimeError("python-docx is unavailable and pandoc is not installed")
    source = md_path
    tmp_name: str | None = None
    original_text = md_path.read_text(encoding="utf-8")
    text = original_text
    text = re.sub(r"```text\s*\nSTOP_FOR_USER\n.*?```", "", text, flags=re.S)
    text = re.sub(r"<!--[^>]*截图[^>]*-->", "【截图预留：请在此处插入当前功能页面或操作结果截图。】", text)
    text = strip_markdown_links(text)
    if code_mode:
        text = re.sub(r"(?=^##\s+第\s*\d+\s*页)", r"\n\\newpage\n", text, flags=re.M)
    if code_mode or "STOP_FOR_USER" in original_text:
        with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
            tmp.write(text)
            tmp_name = tmp.name
        source = Path(tmp_name)
    try:
        subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", str(source), "-o", str(out_path)], check=True)
    finally:
        if tmp_name:
            Path(tmp_name).unlink(missing_ok=True)


def build_code_docx(
    md_path: Path, out_path: Path, software_name: str, version: str, start_page: int = 1
) -> int:
    """返回本文件重排后的页数，供调用方给下一文件续接页码。"""
    if DOCX_AVAILABLE:
        n = build_code_docx_python(md_path, out_path, software_name, version, start_page)
    else:
        n = build_code_docx_ooxml(md_path, out_path, software_name, version, start_page)
    normalize_docx_text_color(out_path)
    return n


def build_manual_docx(md_path: Path, out_path: Path, base_dir: Path, software_name: str, version: str) -> None:
    if DOCX_AVAILABLE:
        build_manual_docx_python(md_path, out_path, base_dir, software_name, version)
    else:
        build_with_pandoc(md_path, out_path, code_mode=False)
        add_header_to_existing_docx(out_path, f"{software_name} {version}")
    normalize_docx_text_color(out_path)


def run_command(command: list[str], cwd: Path | None = None, timeout: int = 60) -> tuple[int, str]:
    try:
        completed = subprocess.run(command, cwd=cwd, text=True, capture_output=True, timeout=timeout)
        return completed.returncode, (completed.stdout + completed.stderr).strip()
    except Exception as exc:
        return 99, str(exc)


def docx_checks(skill_dir: Path, outputs: list[Path]) -> list[str]:
    notes: list[str] = []
    env_script = skill_dir / "vendor/docx-toolkit/scripts/env_check.sh"
    preview_script = skill_dir / "vendor/docx-toolkit/scripts/docx_preview.sh"
    if env_script.exists():
        code, output = run_command(["bash", str(env_script)], cwd=env_script.parent.parent, timeout=30)
        status = "READY" if code == 0 else "NOT READY"
        first_lines = "\n".join(output.splitlines()[:12])
        notes.append(f"DOCX env: {status}\n\n```text\n{first_lines}\n```")
    else:
        notes.append("DOCX env: vendor script missing")

    if preview_script.exists():
        for out in outputs:
            code, output = run_command(["bash", str(preview_script), str(out)], timeout=45)
            first_lines = "\n".join(output.splitlines()[:8])
            notes.append(f"Preview {out.name}: exit={code}\n\n```text\n{first_lines}\n```")
    return notes


def build_all(workdir: Path, software_name: str, version: str, skip_preview: bool) -> dict[str, Any]:
    workdir = ensure_dir(workdir)
    draft_dir = workdir / "草稿"
    final_dir = ensure_dir(workdir / "正式资料")
    app_name = application_software_name(draft_dir)
    app_version = application_version(draft_dir)
    final_software_name = app_name or software_name
    final_version = app_version or version
    safe_name = safe_filename(final_software_name)
    outputs: list[Path] = []
    warnings: list[str] = []
    if app_name and app_name != software_name:
        warnings.append(f"命令参数软件名称为 {software_name}，正式资料已按申请表信息软件名称 {app_name} 生成")
    if app_version and app_version != version:
        warnings.append(f"命令参数版本号为 {version}，正式资料已按申请表信息版本号 {app_version} 生成")
    screenshot_confirmation = read_json_if_exists(workdir / "截图方式确认.json")
    screenshot_method = screenshot_confirmation.get("screenshot_method")
    screenshot_manifest = workdir / "截图/截图清单.json"
    if screenshot_method == "skip":
        warnings.append("用户选择暂不截图；操作手册已保留截图预留位置")
    elif screenshot_method and not screenshot_manifest.exists():
        warnings.append("操作手册截图未生成或未插入；操作手册应保留截图预留位置")
    elif screenshot_manifest.exists():
        screenshots = read_json_if_exists(screenshot_manifest).get("screenshots") or []
        if not screenshots:
            warnings.append("操作手册截图清单为空；操作手册应保留截图预留位置")

    # ⛔ 先生成代码 docx 拿到重排后的真实总页数，再写申请表（回填页数），保证两者一致。
    code_specs = [
        ("代码-前30页.md", f"{safe_name}-代码(前30页).docx"),
        ("代码-后30页.md", f"{safe_name}-代码(后30页).docx"),
        ("代码-全部.md", f"{safe_name}-代码(全部).docx"),
    ]
    # 跨文件连续编号：前30页从 1 起，后30页接前一文件重排后的实际末页+1
    # （重排会改变页数，不能再写死 31；按上一文件返回的页数动态续接）
    next_start = 1
    total_code_pages = 0
    for md_name, docx_name in code_specs:
        md_path = draft_dir / md_name
        if md_path.exists():
            out_path = final_dir / docx_name
            n_pages = build_code_docx(md_path, out_path, final_software_name, final_version, start_page=next_start)
            next_start += n_pages
            total_code_pages += n_pages
            outputs.append(out_path)

    # 用重排后的真实总页数回填申请表 ➤页数（放在代码生成之后）
    app_txt, app_warnings = write_application_txt(draft_dir, final_dir, total_pages=total_code_pages)
    if app_txt:
        outputs.append(app_txt)
    warnings.extend(app_warnings)

    manual_md = draft_dir / "操作手册.md"
    if manual_md.exists():
        manual_out = final_dir / f"{safe_name}_操作手册.docx"
        manual_source = manual_md
        tmp_manual: Path | None = None
        if app_name and app_name != software_name:
            text = manual_md.read_text(encoding="utf-8").replace(software_name, app_name)
            with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
                tmp.write(text)
                tmp_manual = Path(tmp.name)
            manual_source = tmp_manual
        try:
            build_manual_docx(manual_source, manual_out, draft_dir, final_software_name, final_version)
        finally:
            if tmp_manual:
                tmp_manual.unlink(missing_ok=True)
        outputs.append(manual_out)
    else:
        warnings.append("缺少草稿/操作手册.md")

    skill_dir = Path(__file__).resolve().parents[1]
    notes = [] if skip_preview else docx_checks(skill_dir, [p for p in outputs if p.suffix.lower() == ".docx"])
    report = write_report(final_dir, outputs, warnings, notes)
    return {"outputs": [str(p) for p in outputs], "warnings": warnings, "report": str(report)}


def write_report(workdir: Path, outputs: list[Path], warnings: list[str], notes: list[str]) -> Path:
    report = workdir / "生成报告.md"
    lines = ["# 生成报告", "", "## 输出文件", ""]
    for path in outputs:
        size = path.stat().st_size if path.exists() else 0
        lines.append(f"- `{path.name}` ({size} bytes)")
    lines.extend(["", "## 警告", ""])
    if warnings:
        lines.extend(f"- {warning}" for warning in warnings)
    else:
        lines.append("- 无")
    lines.extend(["", "## DOCX 校验", ""])
    if notes:
        lines.extend(notes)
    else:
        lines.append("- 已跳过预览校验")
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workdir", default="软件著作权申请资料")
    parser.add_argument("--software-name", required=True)
    parser.add_argument("--version", default="V1.0")
    parser.add_argument("--skip-preview", action="store_true")
    args = parser.parse_args()

    workdir = Path(args.workdir)
    issues = confirmation_issues(workdir)
    if issues:
        print("STOP_FOR_USER")
        print("NEXT_ACTION: 正式 Word/TXT 生成前必须完成以下确认：")
        for issue in issues:
            print(f"- {issue}")
        raise SystemExit(2)

    result = build_all(workdir, args.software_name, args.version, args.skip_preview)
    print(f"OK final materials: {Path(args.workdir) / '正式资料'}")
    for output in result["outputs"]:
        print(output)
    if result["warnings"]:
        print("Warnings:")
        for warning in result["warnings"]:
            print(f"- {warning}")
    print(f"Report: {result['report']}")


if __name__ == "__main__":
    main()
